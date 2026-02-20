from datetime import datetime
import io
from io import BytesIO # Safe import

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go # Added for go.Figure
import streamlit as st
import numpy as np # For numerical checks


from professional_report import PremiumPlotivaReport

# Try to import new PPTX util, handle if missing for robustness
try:
    from utils_pptx import create_premium_pptx
except ImportError:
    create_premium_pptx = None

# Try to import new PPTX util, handle if missing for robustness
try:
    from utils_pptx import create_premium_pptx
except ImportError:
    create_premium_pptx = None

# --- HELPER: Recreate Chart from Config ---
def recreate_chart(widget_conf, df):
    """Rebuilds a Plotly figure based on the stored configuration."""
    try:
        w_type = widget_conf.get('type')
        if w_type == 'Chart':
            c_type = widget_conf.get('chart_type')
            x = widget_conf.get('x')
            y = widget_conf.get('y')
            color = widget_conf.get('color')
            
            if x and y and x in df.columns and y in df.columns:
                if c_type == "Bar":
                    return px.bar(df, x=x, y=y, color=color, template="plotly_white", title=f"{y} by {x}")
                elif c_type == "Line":
                    return px.line(df, x=x, y=y, color=color, template="plotly_white", title=f"{y} over {x}")
                elif c_type == "Scatter":
                    return px.scatter(df, x=x, y=y, color=color, template="plotly_white", title=f"{y} vs {x}")
                elif c_type == "Pie":
                    return px.pie(df, names=x, values=y, template="plotly_white", title=f"{y} Distribution")
    except Exception:
        pass
    return None

def generate_chart_insight(fig, df):
    """Auto-generate meaningful caption from chart data"""
    try:
        chart_type = fig.data[0].type
        if chart_type == 'scatter':
            return f"Scatter plot analysis showing relationship between variables."
        elif chart_type == 'box':
            return "Distribution analysis showing median and variance across categories"
        elif chart_type == 'histogram':
            return "Frequency distribution showing data spread and central tendency"
        elif chart_type == 'bar':
             return "Categorical comparison highlight key performance differences"
        return "Data visualization showing distribution and patterns"
    except:
        return "Visual analysis of key metrics"

def render_report_tab():
    """Refer the Premium Report Builder UI"""
    
    st.markdown("""
    <div class="premium-card animate-enter" style="margin-bottom: 2rem;">
        <div style="display: flex; justify-content: space-between; align-items: center;">
            <div>
                <h2 style="font-family: 'Playfair Display', serif; font-size: 2.2rem; margin: 0; color: var(--text-main);">Executive Reporting Suite</h2>
                <p style="color: var(--text-secondary); margin-top: 0.5rem; font-family: 'Inter', sans-serif;">
                    Craft high-impact, C-suite ready PDF presentations directly from your analysis and dashboards.
                </p>
            </div>
            <div style="background: rgba(var(--accent-rgb), 0.1); padding: 0.5rem 1rem; border-radius: 20px; border: 1px solid var(--accent);">
                <span style="color: var(--accent); font-weight: 600; font-size: 0.9rem;">Enterprise Grade 💎</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Main columns
    col_conf, col_build = st.columns([1, 2.5], gap="large")

    with col_conf:
        st.markdown("### ⚙️ Configuration")
        with st.expander("📄 Document Meta", expanded=True):
            r_title = st.text_input("Report Title", "Strategic Performance Review")
            r_subtitle = st.text_input("Subtitle", "Data-Driven Insights & Executive Summary")
            r_author = st.text_input("Author Name", "Data Science Team")
            r_company = st.text_input("Company", "Plotiva Enterprise")
        
        with st.expander("🎨 Visual Style", expanded=True):
            st.selectbox("Layout Template", ["Standard Report", "Executive Brief"])
            color_scheme = st.selectbox("Color Theme", ["corporate_blue", "modern_teal", "professional_gray"])
            page_size = st.selectbox("Paper Size", ["letter", "a4"])
            include_methodology = st.checkbox("Include Methodology Appendix", value=True)

    with col_build:
        st.markdown("### 📝 Content Builder")
        
        tab_sum, tab_dash, tab_cust, tab_vis, tab_find, tab_rec = st.tabs([
            "📝 Executive Brief", "💎 Key Metrics", "📊 Studio Boards", "📉 Deep Dive", "🧠 Findings", "🎯 Strategy"
        ])
        
        # 0. Executive Summary
        with tab_sum:
            st.subheader("Executive Summary")
            r_summary = st.text_area("Summary Narrative", 
                "This strategic assessment evaluates current performance trends to identify opportunities for operational optimization and revenue growth...", height=150)
            st.markdown("#### Key Highlights")
            col_h1, col_h2 = st.columns(2)
            h1 = col_h1.text_input("Highlight 1", "Revenue increased 15% YoY")
            h2 = col_h1.text_input("Highlight 2", "CAC decreased 8%")
            h3 = col_h2.text_input("Highlight 3", "Regional variance identified")
            h4 = col_h2.text_input("Highlight 4", "Operational efficiency up 12%")
            highlights = [h for h in [h1, h2, h3, h4] if h]
        
        # 1. Main KPIs
        with tab_dash:
            st.subheader("Report Main Dashboard")
            st.caption("These appear at the very beginning of the report.")
            num_kpis = st.number_input("Number of KP Cards", 1, 6, 4)
            kpi_data = []
            cols = st.columns(2)
            for i in range(num_kpis):
                with cols[i%2]:
                    with st.container(border=True):
                        kn = st.text_input(f"Name {i}", f"Metric {i+1}", key=f"kpi_name_{i}")
                        kv = st.text_input(f"Value {i}", "$0", key=f"kpi_val_{i}")
                        kc = st.text_input(f"Change {i}", "+0%", key=f"kpi_change_{i}")
                        kpi_data.append({'name': kn, 'value': kv, 'change': kc, 'status': 'success'})

        # 2. Custom Dashboards (NEW)
        with tab_cust:
            st.subheader("Include Studio Dashboards")
            if 'custom_dashboards' in st.session_state and st.session_state.custom_dashboards:
                avail_dash = list(st.session_state.custom_dashboards.keys())
                dashboards_to_include = st.multiselect("Select Dashboards to Append", avail_dash, default=avail_dash)
                st.info(f"{len(dashboards_to_include)} dashboard(s) selected for export.")
            else:
                dashboards_to_include = []
                st.warning("No custom dashboards found. Create one in the Dashboard Builder tab.")

        # 3. Visuals
        with tab_vis:
            st.subheader("Individual Analysis Plots")
            use_demo = st.checkbox("Include Demo Plot if none selected", value=True)

        # 4. Findings
        with tab_find:
            st.subheader("Key Findings")
            findings = []
            # Simplified for brevity in this full-file replacement, logic remains same as before
            if st.checkbox("Add Example Finding", value=True):
                 findings.append({'text': "<b>Revenue Acceleration</b><br/>Q3 performance exceeds projections by 15%.", 'type': 'success'})

        # 5. Actions
        with tab_rec:
            st.subheader("Recommendations")
            recs = []
            if st.checkbox("Add Example Recommendation", value=True):
                recs.append({'title': "Scale Acquisition Channels", 'description': "Allocate additional 20% budget to high-performing campaigns.", 'priority': 'high'})

    # --- GENERATION ---
    st.markdown("---")
    c_gen, c_status = st.columns([1, 2])
    
    # Common Data
    df_main = st.session_state.get('df_main', st.session_state.get('working_data', pd.DataFrame()))

    with c_gen:
        if st.button("🎯 Generate Premium Report", type="primary", use_container_width=True):
            progress = st.progress(0, text="Initializing...")
            
            try:
                # ---------------- PDF GENERATION ----------------
                config = {'title': r_title, 'subtitle': r_subtitle, 'author': r_author, 'company': r_company, 'color_scheme': color_scheme, 'page_size': page_size}
                report = PremiumPlotivaReport(config)
                report.add_cover_page()
                report.add_executive_summary(r_summary, highlights)
                report.add_page_break()
                report.add_executive_dashboard(kpi_data)
                
                # Custom Dashboards Integration
                if dashboards_to_include and 'custom_dashboards' in st.session_state:
                    progress.progress(40, text="📊 Rendering Custom Dashboards...")
                    for d_name in dashboards_to_include:
                        d_data = st.session_state.custom_dashboards[d_name]
                        widgets = d_data.get('widgets', {})
                        
                        report.add_colored_section(f"Dashboard: {d_name}")
                        
                        # Render widgets
                        for w_key, w_conf in widgets.items():
                            w_type = w_conf.get('type')
                            if w_type == 'Chart':
                                fig = recreate_chart(w_conf, df_main)
                                if fig:
                                    report.add_plot_with_caption(fig, title=f"{d_name} - Analysis", caption="Dashboard Visualization")
                            elif w_type == 'KPI':
                                # Add simple text for KPI
                                val_col = w_conf.get('val_col')
                                label = w_conf.get('label', 'Metric')
                                val = df_main[val_col].sum() if (val_col and val_col in df_main.columns) else 0
                                report.add_callout_box(f"<b>{label}</b>: {val:,.0f}", 'info')
                        
                        report.add_page_break()

                # Visuals
                if 'saved_figures' in st.session_state and st.session_state.saved_figures:
                    report.add_colored_section("Deep Dive Analysis")
                    for plot in st.session_state.saved_figures[:5]:
                        report.add_plot_with_caption(plot['figure'], title=plot.get('name'), caption=plot.get('caption'))

                if findings:
                    report.add_colored_section("Strategic Insights")
                    for f in findings: report.add_callout_box(f['text'], f['type'])
                
                if recs:
                    report.add_recommendations(recs)
                
                # Methodology Section
                if include_methodology:
                    report.add_page_break()
                    report.add_colored_section("Methodology & Definitions")
                    report.add_text("This report utilizes advanced statistical methods and machine learning algorithms to derive insights. Key definitions include:")
                    report.add_text("<b>Data Integrity</b>: All checks were performed using Plotiva's proprietary quality engine. Missing values were imputed using median strategies to preserve distribution integrity.")
                    report.add_text("<b>Significance Testing</b>: P-values < 0.05 are considered statistically significant for all hypothesis tests presented.")
                    report.add_text("<b>Outlier Detection</b>: Interquartile Range (IQR) method used to identify anomalies.")
                
                pdf_bytes = report.generate()
                st.session_state['last_pdf'] = pdf_bytes
                progress.progress(100, text="Done!")
                st.balloons()
                
            except Exception as e:
                st.error(f"PDF Gen Error: {str(e)}")

    with c_status:
        if 'last_pdf' in st.session_state:
            st.success("✅ PDF Ready")
            st.download_button("📥 Download PDF", st.session_state['last_pdf'], "Report.pdf", "application/pdf", use_container_width=True)

        # ---------------- WORD EXPORT ----------------
        if st.button("📝 Export to Word (Beta)", type="secondary", use_container_width=True):
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Title
                for i in range(5): doc.add_paragraph()
                t = doc.add_heading(r_title, 0)
                t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(r_subtitle).alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_page_break()
                
                # Summary
                doc.add_heading('Executive Summary', 1)
                doc.add_paragraph(r_summary)
                
                # Custom Dashboards
                if dashboards_to_include and 'custom_dashboards' in st.session_state:
                    doc.add_heading('Executive Dashboards', 1)
                    
                    for d_name in dashboards_to_include:
                        doc.add_heading(d_name, 2)
                        d_data = st.session_state.custom_dashboards[d_name]
                        
                        for w_key, w_conf in d_data.get('widgets', {}).items():
                            if w_conf.get('type') == 'Chart':
                                fig = recreate_chart(w_conf, df_main)
                                if fig:
                                    img_bytes = fig.to_image(format="png", width=600, height=350, scale=1.5)
                                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
                                    doc.add_paragraph()
                            elif w_conf.get('type') == 'KPI':
                                val_col = w_conf.get('val_col')
                                label = w_conf.get('label')
                                val = df_main[val_col].sum() if (val_col and val_col in df_main.columns) else 0
                                p = doc.add_paragraph()
                                p.add_run(f"{label}: ").bold = True
                                p.add_run(f"{val:,.0f}")

                # Save
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button("📥 Download Word Doc", buffer, "Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            except Exception as e:
                st.error(f"Word Export Error: {str(e)}")

        # ---------------- EXCEL EXPORT ----------------
        if st.button("📊 Export Data to Excel", type="secondary", use_container_width=True):
             try:
                from openpyxl.styles import PatternFill, Font, Alignment
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    if 'working_data' in st.session_state:
                         st.session_state.working_data.to_excel(writer, sheet_name='Raw Data', index=False)
                    
                    # Dashboard Data Summary
                    if dashboards_to_include:
                        dash_summary = []
                        for d_name in dashboards_to_include:
                            widgets = st.session_state.custom_dashboards[d_name].get('widgets', {})
                            for w_key, w_conf in widgets.items():
                                if w_conf.get('type') == 'KPI':
                                    val_col = w_conf.get('val_col')
                                    val = df_main[val_col].sum() if (val_col and val_col in df_main.columns) else 0
                                    dash_summary.append({'Dashboard': d_name, 'Metric': w_conf.get('label'), 'Value': val})
                        
                        if dash_summary:
                            pd.DataFrame(dash_summary).to_excel(writer, sheet_name='Dashboard Metrics', index=False)

                    # Styling
                    workbook = writer.book
                    teal_fill = PatternFill(start_color="0F766E", end_color="0F766E", fill_type="solid")
                    white_font = Font(color="FFFFFF", bold=True)
                    for sheet in writer.sheets.values():
                        for cell in sheet[1]:
                            cell.fill = teal_fill
                            cell.font = white_font
                            cell.alignment = Alignment(horizontal='center')

                buffer.seek(0)
                st.download_button("📥 Download Excel", buffer, "Data.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
             except Exception as e:
                 st.error(f"Excel Export Error: {str(e)}")

        # ---------------- PPTX EXPORT ----------------
        # ---------------- PPTX EXPORT ----------------
        if st.button("📽️ Export to PowerPoint", type="secondary", use_container_width=True):
            if create_premium_pptx is None:
                st.error("PPTX Utility not found. Please verify installation.")
            else:
                try:
                    # Prepare Data for PPTX
                    pptx_config = {
                        'title': r_title,
                        'subtitle': r_subtitle,
                        'company': r_company,
                        'summary': r_summary
                    }
                    
                    # Collect Charts
                    charts = []
                    if 'saved_figures' in st.session_state:
                        charts = st.session_state.saved_figures
                    
                    # Generate
                    prs = create_premium_pptx(pptx_config, df_main, kpi_data, findings, recs, charts)
                    
                    # Save
                    buffer = io.BytesIO()
                    prs.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        "📥 Download Presentation", 
                        buffer, 
                        "Presentation.pptx", 
                        "application/vnd.openxmlformats-officedocument.presentationml.presentation", 
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"PPTX Export Failed: {str(e)}")


def generate_chart_insight(fig, df):
    """Auto-generate meaningful caption from chart data"""
    try:
        chart_type = fig.data[0].type
        
        if chart_type == 'scatter':
            # Calculate correlation
            try:
                fig.data[0].x
                fig.data[0].y
                # If x/y are arrays, we can't easily correlate without DF reference or rebuilding series
                # But typically plotly names match df columns if generated via px
                # Let's try basic text
                return f"Scatter plot analysis showing relationship between {fig.layout.xaxis.title.text} and {fig.layout.yaxis.title.text}"
            except:
                return "Scatter plot analysis of variable relationships"
        
        elif chart_type == 'box':
            return "Distribution analysis showing median and variance across categories"
        
        elif chart_type == 'histogram':
            return "Frequency distribution showing data spread and central tendency"
            
        elif chart_type == 'bar':
             return "Categorical comparison highlight key performance differences"
             
        return "Data visualization showing distribution and patterns"
    except:
        return "Visual analysis of key metrics"

def render_report_tab():
    """Refer the Premium Report Builder UI"""
    
    st.markdown("""
    <div class="premium-card animate-enter" style="margin-bottom: 2rem;">
        <div style="display: flex; justify-content: space-between; align-items: center;">
            <div>
                <h2 style="font-family: 'Playfair Display', serif; font-size: 2.2rem; margin: 0; color: var(--text-main);">Executive Reporting Suite</h2>
                <p style="color: var(--text-secondary); margin-top: 0.5rem; font-family: 'Inter', sans-serif;">
                    Craft high-impact, C-suite ready PDF presentations directly from your analysis.
                </p>
            </div>
            <div style="background: rgba(var(--accent-rgb), 0.1); padding: 0.5rem 1rem; border-radius: 20px; border: 1px solid var(--accent);">
                <span style="color: var(--accent); font-weight: 600; font-size: 0.9rem;">Enterprise Grade 💎</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Main columns: Configuration (Left) vs Content Builder (Center/Right)
    col_conf, col_build = st.columns([1, 2.5], gap="large")

    with col_conf:
        st.markdown("### ⚙️ Configuration")
        
        with st.expander("📄 Document Meta", expanded=True):
            r_title = st.text_input("Report Title", "Business Strategy Report")
            r_subtitle = st.text_input("Subtitle", "Data-Driven Insights & Executive Summary")
            r_author = st.text_input("Author Name", "Data Science Team")
            r_company = st.text_input("Company", "Plotiva Enterprise")
        
        with st.expander("🎨 Visual Style", expanded=True):
            st.selectbox("Layout Template", ["Standard Report", "Executive Brief", "Technical Analysis"])
            color_scheme = st.selectbox("Color Theme", ["corporate_blue", "modern_teal", "professional_gray"])
            page_size = st.selectbox("Paper Size", ["letter", "a4"])
        
        with st.expander("📊 Chart Styling"):
            st.caption("Apply consistent palette to all charts")
            st.multiselect("Active Palette", ["#1f77b4", "#ff7f0e", "#2ca02c", "#d62728"], default=["#1f77b4", "#ff7f0e", "#2ca02c"], disabled=True)

    with col_build:
        st.markdown("### 📝 Content Builder")
        
        # Using Tabs for different section types
        tab_sum, tab_dash, tab_vis, tab_find, tab_rec, tab_app = st.tabs([
            "📑 Executive Summary", "📊 Dashboard", "📈 Visuals", "🧠 Findings", "🚀 Recommendations", "📎 Appendix"
        ])
        
        # 0. Executive Summary (New)
        with tab_sum:
            st.markdown("<div class='control-panel-card'>", unsafe_allow_html=True)
            st.subheader("Executive Summary")
            st.caption("Craft a high-level overview for stakeholders.")
            
            # Pro Tips / Templates
            with st.expander("💡 Pro Tip: Content Templates"):
                st.markdown("**Executive Summary Template:**")
                st.code("This analysis examines [dataset description] covering [time period]. Key objectives include [1) objective, 2) objective]. The findings reveal [major insight] with significant implications for [business area].")
                
                st.markdown("**Insight Examples:**")
                st.text("• Strong positive correlation (r=0.85) observed between [var1] and [var2]")
                st.text("• Regional analysis reveals [Region A] outperforming by 15%, driven by [factor]")
                st.text("• Temporal trends indicate seasonality with peak performance during [period]")
            
            r_summary = st.text_area("Summary Narrative", 
                "This report provides a comprehensive analysis of key performance metrics and underlying data trends. The findings highlight significant opportunities for optimization and strategic growth.",
                height=150)
            
            st.markdown("#### Key Highlights")
            col_h1, col_h2 = st.columns(2)
            highlights = []
            with col_h1:
                h1 = st.text_input("Highlight 1", "Revenue increased 15% YoY driven by enterprise segment growth")
                h2 = st.text_input("Highlight 2", "Customer acquisition costs decreased 8% through optimization")
            with col_h2:
                h3 = st.text_input("Highlight 3", "Regional performance variance identified (North +22%, South +8%)")
                h4 = st.text_input("Highlight 4", "Operational efficiency improved by 12% in Q3")
            
            highlights = [h for h in [h1, h2, h3, h4] if h]
            st.markdown("</div>", unsafe_allow_html=True)
        
        # 1. Dashboard
        with tab_dash:
            st.markdown("<div class='control-panel-card'>", unsafe_allow_html=True)
            st.subheader("Executive KPI Dashboard")
            st.caption("Define up to 6 high-level metrics for the opening dashboard page.")
            
            num_kpis = st.number_input("Number of KPI Cards", 1, 6, 4)
            
            kpi_data = []
            # Grid layout for inputs
            cols = st.columns(2)
            for i in range(num_kpis):
                with cols[i%2]:
                    with st.container(border=True):
                        st.markdown(f"**KPI #{i+1}**")
                        kn = st.text_input(f"Name", st.session_state.get(f"kpi_name_{i}", f"Metric {i+1}"), key=f"kpi_name_{i}")
                        kv = st.text_input(f"Value", st.session_state.get(f"kpi_val_{i}", "$0"), key=f"kpi_val_{i}")
                        kc = st.text_input(f"Change", st.session_state.get(f"kpi_change_{i}", "+0% YoY"), key=f"kpi_change_{i}")
                        ks = st.selectbox(f"Status", ["success", "warning", "danger"], index=["success", "warning", "danger"].index(st.session_state.get(f"kpi_status_{i}", "success")), key=f"kpi_status_{i}")
                        kpi_data.append({'name': kn, 'value': kv, 'change': kc, 'status': ks})
            st.markdown("</div>", unsafe_allow_html=True)
            
        # 2. Visuals
        with tab_vis:
            st.markdown("<div class='control-panel-card'>", unsafe_allow_html=True)
            st.subheader("Analysis Plots")
            st.info("Select plots from your current session to include in the report.")
            
            # Integration with existing session state plots
            if 'saved_figures' in st.session_state and st.session_state.saved_figures:
                 st.success(f"Found {len(st.session_state.saved_figures)} saved plots available for reporting.")
            else:
                 st.warning("No saved plots found in this session. Generate complex plots in the Exploration tab, or use the demo plotting below.")
                 
            # Demo Plot Logic
            use_demo = st.checkbox("Include Demo Analysis Plot", value=True)
            
            st.markdown("</div>", unsafe_allow_html=True)

        # 3. Findings
        with tab_find:
            st.markdown("<div class='control-panel-card'>", unsafe_allow_html=True)
            st.subheader("Key Findings & Insights")
            
            findings = []
            
            use_auto_insights = st.checkbox("✨ Auto-populate from Insights Engine", value=True, help="Use AI-driven statistical analysis to generate findings automatically.")
            
            if not use_auto_insights:
                num_find = st.number_input("Number of Findings", 1, 5, 3)
                for i in range(num_find):
                    with st.expander(f"Finding #{i+1}", expanded=(i==0)):
                        ft = st.text_input(f"Title", f"Strategic Insight {i+1}", key=f"f_{i}_t")
                        
                        # Pre-fill smart defaults suggestion
                        insight_placeholder = "Example: Regional analysis reveals North outperforming by 22%, driven by 3 major enterprise contracts in Q4."
                        fd = st.text_area(f"Description", 
                                         value="Key observation regarding the data..." if i > 0 else "Analysis indicates a significant trend in the primary metrics...",
                                         placeholder=insight_placeholder,
                                         key=f"f_{i}_d")
                        fs = st.selectbox(f"Type", ["info", "success", "warning", "danger"], key=f"f_{i}_s")
                        findings.append({'text': f"<b>{ft}</b><br/>{fd}", 'type': fs})
            st.markdown("</div>", unsafe_allow_html=True)

        # 4. Recommendations
        with tab_rec:
            st.markdown("<div class='control-panel-card'>", unsafe_allow_html=True)
            st.subheader("Strategic Recommendations")
            
            num_rec = st.number_input("Number of Action Items", 1, 5, 3)
            recs = []
            for i in range(num_rec):
                with st.expander(f"Recommendation #{i+1}", expanded=(i==0)):
                    rt = st.text_input(f"Action Title", f"Strategic Initiative {i+1}", key=f"r_{i}_t")
                    rd = st.text_area(f"Details", "Proposed course of action based on analysis...", key=f"r_{i}_d")
                    rp = st.radio(f"Priority", ["low", "medium", "high"], index=2, horizontal=True, key=f"r_{i}_p")
                    recs.append({'title': rt, 'description': rd, 'priority': rp})
            st.markdown("</div>", unsafe_allow_html=True)

        # 5. Appendix
        with tab_app:
            st.markdown("<div class='control-panel-card'>", unsafe_allow_html=True)
            st.subheader("Technical Appendix")
            
            inc_meth = st.checkbox("Include Methodology Section", value=True)
            methods = []
            if inc_meth:
                m_text = st.text_area("Methodology Description", 
                                      "1. Sourcing: Data loaded from local ingestion.\n2. Cleaning: Nulls handled via imputation.\n3. Analysis: Standard statistical features used.",
                                      height=100)
                methods = m_text.split('\n')
                
            st.markdown("</div>", unsafe_allow_html=True)

    # Global Action Bar
    st.markdown("---")
    st.header("🎯 Finalize Report")

    # Check state
    is_generated = 'last_pdf' in st.session_state and st.session_state['last_pdf'] is not None

    if not is_generated:
        # --- STATE 1: GENERATION PHASE ---
        col_gen_main, col_gen_help = st.columns([1, 1.5], gap="large")
        
        with col_gen_main:
            if st.button("🎯 Generate Premium Report", type="primary", use_container_width=True):
                
                # 1. Validation
                if not kpi_data or all(k['value'] == '$0' for k in kpi_data):
                    st.error("⚠️ All KPIs are default ($0). Please enter real values!")
                    st.stop()
                
                if 'saved_figures' not in st.session_state and not use_demo:
                     st.warning("No plots selected. Go to Visuals tab to simple/demo charts.")

                # 2. Initialization
                progress = st.progress(0, text="Initializing...")
                
                try:
                    # Config
                    config = {
                        'title': r_title,
                        'subtitle': r_subtitle,
                        'author': r_author,
                        'company': r_company,
                        'color_scheme': color_scheme,
                        'page_size': page_size
                    }
                    
                    progress.progress(10, text="📄 Building report structure...")
                    report = PremiumPlotivaReport(config)
                    
                    progress.progress(20, text="📊 Adding cover page...")
                    report.add_cover_page()
                    
                    # EXECUTIVE SUMMARY
                    progress.progress(25, text="📑 Adding executive summary...")
                    report.add_executive_summary(r_summary, highlights)
                    report.add_page_break()

                    progress.progress(30, text="📈 Creating dashboard...")
                    report.add_executive_dashboard(kpi_data)
                    
                    # DATASET OVERVIEW
                    if 'df_main' in st.session_state and st.session_state.df_main is not None:
                         report.add_comprehensive_data_overview(st.session_state.df_main)
                         report.add_page_break()
                    
                    progress.progress(50, text="🎨 Adding insights...")
                    
                    # Auto-generate findings if requested
                    if use_auto_insights:
                        try:
                            # Ensure we have data
                            if 'df_main' in st.session_state:
                                from premium_insights import generate_executive_insights
                                auto_insights = generate_executive_insights(st.session_state.df_main)
                                findings = [{'text': i['text'], 'type': 'success' if i['priority'] == 'high' else 'info'} for i in auto_insights]
                            else:
                                st.warning("Data not available for automated insights.")
                        except ImportError:
                            pass
                    
                    if findings:
                        report.add_colored_section("Executive Insights")
                        for finding in findings:
                            report.add_callout_box(finding['text'], finding['type'])
                    
                    progress.progress(60, text="📊 Rendering visualizations...")
                    report.add_colored_section("Visual Analysis")
                    
                    report.add_text("The following visualizations provide detailed analysis of key performance drivers and underlying data patterns.")

                    # DIAGNOSTIC
                    report.add_colored_section("Diagnostic Analysis")
                    report.add_text("This section explores correlations and variance to understand underlying drivers.")
                    
                    # PREDICTIVE
                    if 'ml_results' in st.session_state and st.session_state.ml_results.get('history'):
                        report.add_colored_section("Predictive Intelligence")
                        best_model = st.session_state.ml_results['history'][-1]
                        report.add_text(f"A {best_model.get('model_name', 'Model')} was trained to predict target outcomes.")
                        metric_txt = ", ".join([f"{k}: {v:.2f}" for k,v in best_model.get('metrics', {}).items() if isinstance(v, (int, float))])
                        report.add_callout_box(f"Model Performance: {metric_txt}", 'success')
                    
                    # VISUALS
                    report.add_colored_section("Visual Analysis")

                    plots_added = 0
                    if 'saved_figures' in st.session_state and st.session_state.saved_figures:
                        for plot in st.session_state.saved_figures[:6]:  # Limit to 6
                            try:
                                caption = plot.get('caption', '')
                                report.add_plot_with_caption(
                                    plot['figure'],
                                    title=plot.get('name', ''),
                                    caption=caption,
                                    size='medium',
                                    df=st.session_state.df_main if 'df_main' in st.session_state else None
                                )
                                plots_added += 1
                            except Exception as e:
                                st.warning(f"Skipped plot '{plot.get('name')}': {str(e)}")
                                
                    if use_demo and plots_added == 0:
                         df_demo = px.data.tips()
                         fig = px.bar(df_demo, x="day", y="total_bill", color="sex", barmode="group")
                         report.add_plot_with_caption(fig, title="Sample Analysis", caption="Comparative analysis of total bill distribution by day and gender")
                    
                    progress.progress(80, text="💡 Adding recommendations...")
                    if recs:
                        report.add_recommendations(recs)
                    
                    progress.progress(90, text="📎 Finalizing...")
                    report.add_appendix()
                    if inc_meth:
                        report.add_methodology_section(methods)
                    
                    # GLOSSARY
                    report.add_colored_section("Glossary of Terms")
                    report.add_text("<b>IQR</b>: Interquartile Range, measuring statistical dispersion.")
                    report.add_text("<b>Correlation (r)</b>: Strength of relationship between variables (-1 to 1).")
                    report.add_text("<b>P-Value</b>: Probability of results occurring by chance.")
                    
                    progress.progress(95, text="🎯 Compiling PDF...")
                    pdf_bytes = report.generate()
                    
                    progress.progress(100, text="✅ Complete!")
                    
                    # DONE
                    st.session_state['last_pdf'] = pdf_bytes
                    st.balloons()
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Generation failed: {str(e)}")
                    st.info("Check if 'kaleido' is installed for image export.")

        with col_gen_help:
            st.info("👈 **Ready to finalize?**\n\nConfigure your sections above, then click 'Generate' to compile your executive report. Download options (PDF, PPTX, Word, Excel) will appear after generation.")

    else:
        # --- STATE 2: DOWNLOAD PHASE ---
        st.success("✅ Report generated successfully!")
        
        st.subheader("📥 Download Options")
        
        d_col1, d_col2, d_col3, d_col4 = st.columns(4)
        
        # 1. PDF
        with d_col1:
            st.download_button(
                label="📄 PDF Report",
                data=st.session_state['last_pdf'],
                file_name=f"Executive_Report_{datetime.now().strftime('%Y%m%d')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
            
        # 2. WORD
        with d_col2:
            if st.button("📝 Word Export", use_container_width=True):
                try:
                    from docx import Document
                    from docx.shared import Inches, Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    from docx.oxml.ns import nsdecls
                    from docx.oxml import parse_xml
                    
                    doc = Document()
                    
                    # --- STYLES ---
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Calibri'
                    font.size = Pt(11)
                    
                    # --- COVER PAGE ---
                    for _ in range(3): doc.add_paragraph()
                    
                    # Title
                    p_title = doc.add_paragraph(r_title)
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_title_run = p_title.runs[0]
                    r_title_run.bold = True
                    r_title_run.font.size = Pt(28)
                    r_title_run.font.color.rgb = RGBColor(15, 118, 110) # Teal
                    
                    # Subtitle
                    p_sub = doc.add_paragraph(r_subtitle)
                    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_sub.runs[0].font.size = Pt(16)
                    p_sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)
                    
                    for _ in range(4): doc.add_paragraph()
                    
                    # Meta
                    p_meta = doc.add_paragraph()
                    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_meta = p_meta.add_run(f"Prepared by: {r_author}\n{r_company}\n{datetime.now().strftime('%B %d, %Y')}")
                    run_meta.font.size = Pt(12)
                    
                    doc.add_page_break()
                    
                    # --- EXECUTIVE SUMMARY ---
                    h1 = doc.add_heading('Executive Summary', level=1)
                    h1.runs[0].font.color.rgb = RGBColor(15, 118, 110)
                    
                    doc.add_paragraph(r_summary)
                    
                    if highlights:
                        doc.add_heading('Key Highlights', level=2)
                        for h in highlights:
                            doc.add_paragraph(h, style='List Bullet')
                            
                    doc.add_paragraph()

                    # --- KPI DASHBOARD ---
                    if kpi_data:
                        doc.add_heading('Performance Snapshot', level=1)
                        # Create a table
                        table = doc.add_table(rows=1, cols=4)
                        table.style = 'Table Grid'
                        table.autofit = True
                        
                        # Header
                        hdr_cells = table.rows[0].cells
                        headers = ['Metric', 'Current Value', 'Change', 'Status']
                        for i, h in enumerate(headers):
                            hdr_cells[i].text = h
                            # Shade header
                            tcPr = hdr_cells[i]._element.get_or_add_tcPr()
                            shd = parse_xml(r'<w:shd {} w:fill="0F766E"/>'.format(nsdecls('w')))
                            tcPr.append(shd)
                            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                            hdr_cells[i].paragraphs[0].runs[0].bold = True
                        
                        # Data
                        for kpi in kpi_data:
                            row_cells = table.add_row().cells
                            row_cells[0].text = kpi.get('name', 'N/A')
                            row_cells[1].text = str(kpi.get('value', '0'))
                            row_cells[2].text = str(kpi.get('change', '-'))
                            row_cells[3].text = kpi.get('status', 'info').upper()
                        
                        doc.add_paragraph()

                    # --- VISUAL ANALYSIS ---
                    if 'saved_figures' in st.session_state and st.session_state.saved_figures:
                        doc.add_heading('Visual Analysis', level=1)
                        doc.add_paragraph("Comprehensive data visualization based on current analysis findings.")
                        
                        for i, plot in enumerate(st.session_state.saved_figures[:6]):
                            try:
                                fig = plot['figure']
                                img_bytes = fig.to_image(format="png", width=800, height=500, scale=1.5)
                                img_stream = BytesIO(img_bytes)
                                
                                doc.add_heading(plot.get('name', f'Chart {i+1}'), level=2)
                                doc.add_picture(img_stream, width=Inches(6.2))
                                
                                # Caption box
                                cap = doc.add_paragraph()
                                cap.add_run(f"Insight: ").bold = True
                                cap.add_run(plot.get('caption', 'Analysis of key trends.'))
                                cap.style = 'Caption'
                                doc.add_paragraph()
                            except:
                                pass
                                
                    # --- STRATEGIC INSIGHTS ---
                    if findings:
                        doc.add_heading('Strategic Insights', level=1)
                        for f in findings:
                            p = doc.add_paragraph(style='List Bullet')
                            # Simple HTML strip
                            txt = f['text'].replace('<b>', '').replace('</b>', '').replace('<br/>', ': ')
                            run = p.add_run(txt)
                            
                    # --- RECOMMENDATIONS ---
                    if recs:
                        doc.add_heading('Recommendations', level=1)
                        for r in recs:
                            p = doc.add_paragraph()
                            p.add_run(f"• {r['title']}").bold = True
                            p.add_run(f" ({r['priority'].upper()} Priority)\n")
                            p.add_run(f"  {r['description']}")
                    
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.session_state['last_word'] = buffer
                    st.rerun()
                except Exception as e:
                    st.error(f"Word Error: {str(e)}")
                    
            if 'last_word' in st.session_state:
                 st.download_button("📥 Save .docx", st.session_state['last_word'], "Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_docx_ready", use_container_width=True)

        # 3. EXCEL
        with d_col3:
            if st.button("📊 Excel Data", use_container_width=True):
                try:

                    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
                    
                    buffer = BytesIO()
                    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                        
                        # --- SHEET 1: EXECUTIVE SUMMARY ---
                        # Create a small summary DF
                        summary_data = {
                            'Item': ['Report Title', 'Date', 'Author', 'Total Highlights'],
                            'Value': [r_title, datetime.now().strftime('%Y-%m-%d'), r_author, len(highlights) if highlights else 0]
                        }
                        pd.DataFrame(summary_data).to_excel(writer, sheet_name='Executive Summary', index=False, startrow=1, startcol=1)
                        
                        # --- SHEET 2: KPIS ---
                        if kpi_data:
                            pd.DataFrame(kpi_data).to_excel(writer, sheet_name='KPI Dashboard', index=False)
                        
                        # --- SHEET 3: INSIGHTS ---
                        if findings:
                            f_clean = [{'Insight': f['text'].replace('<b>','').replace('</b>','').replace('<br/>', ': '), 'Type': f['type']} for f in findings]
                            pd.DataFrame(f_clean).to_excel(writer, sheet_name='Insights', index=False)
                            
                        # --- SHEET 4: RECOMMENDATIONS ---
                        if recs:
                            pd.DataFrame(recs).to_excel(writer, sheet_name='Recommendations', index=False)
                            
                        # --- SHEET 5: RAW DATA ---
                        if 'working_data' in st.session_state:
                             st.session_state.working_data.to_excel(writer, sheet_name='Raw Data', index=False)
                        
                        # --- STYLING MACRO ---
                        workbook = writer.book
                        header_fill = PatternFill(start_color="0F766E", end_color="0F766E", fill_type="solid")
                        header_font = Font(color="FFFFFF", bold=True, size=12)
                        
                        for sheet_name in writer.sheets:
                            ws = writer.sheets[sheet_name]
                            
                            # Auto-width & Header Style
                            for column_cells in ws.columns:
                                length = max(len(str(cell.value)) for cell in column_cells)
                                ws.column_dimensions[column_cells[0].column_letter].width = min(length + 2, 60)
                            
                            # Apply to first row
                            for cell in ws[1]:
                                cell.fill = header_fill
                                cell.font = header_font
                                cell.alignment = Alignment(horizontal='center', vertical='center')

                    buffer.seek(0)
                    st.session_state['last_excel'] = buffer
                    st.rerun()
                except Exception as e:
                    st.error(f"Excel Error: {str(e)}")
            
            if 'last_excel' in st.session_state:
                st.download_button("📥 Save .xlsx", st.session_state['last_excel'], "Data.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_xlsx_ready", use_container_width=True)

        # 4. PPTX
        with d_col4:
            if st.button("📽️ PowerPoint", use_container_width=True):
                 try:
                    if create_premium_pptx:
                        pptx_config = {'title': r_title, 'subtitle': r_subtitle, 'company': r_company, 'summary': r_summary}
                        charts = st.session_state.get('saved_figures', [])
                        prs = create_premium_pptx(pptx_config, st.session_state.get('df_main', pd.DataFrame()), kpi_data, findings, recs, charts)
                        buffer = BytesIO()
                        prs.save(buffer)
                        buffer.seek(0)
                        st.session_state['last_pptx'] = buffer
                        st.rerun()
                    else:
                        st.error("PPTX Module missing.")
                 except Exception as e:
                     st.error(f"PPTX Error: {e}")

            if 'last_pptx' in st.session_state:
                st.download_button("📥 Save .pptx", st.session_state['last_pptx'], "Presentation.pptx", "application/vnd.openxmlformats-officedocument.presentationml.presentation", key="dl_pptx_ready")

        st.markdown("---")
        
        # --- ADVANCED EXPORTS ---
        st.subheader("Advanced Data Exports")
        
        col_adv_1, col_adv_2 = st.columns([1, 2])
        
        with col_adv_1:
            if st.button("🐍 Export Analysis Code", use_container_width=True):
                # Generate Python Script
                py_code = f"""
# Plotiva Advanced Analysis Export
# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
# ---------------------------------------------------------
# This script provides a complete template for:
# 1. Loading and Cleaning Data
# 2. Advanced Exploratory Data Analysis (EDA)
# 3. Training a Machine Learning Model (Random Forest)
# 4. Generating SHAP (SHapley Additive exPlanations) Values
# ---------------------------------------------------------

import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestRegressor, RandomForestClassifier
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.impute import SimpleImputer
import shap

# ==========================================
# 1. LOAD DATA & METADATA
# ==========================================
# Replace with your actual file path
DATA_PATH = 'your_data.csv' 
try:
    df = pd.read_csv(DATA_PATH)
    print(f"✅ Data Loaded. Shape: {{df.shape}}")
except FileNotFoundError:
    print("⚠️ File not found. Please update DATA_PATH.")
    # Creating Dummy Data for Demonstration
    df = pd.DataFrame(np.random.randint(0,100,size=(100, 4)), columns=list('ABCD'))
    df['Target'] = np.random.choice([0, 1], size=100)

print("\\n--- Executive Snapshot ---")
{chr(10).join([f'print("Metric: {k["name"]} | Value: {k["value"]}")' for k in kpi_data]) if kpi_data else ''}

# ==========================================
# 2. ADVANCED EDA
# ==========================================
print("\\n--- Running Diagnostic Checks ---")
print(df.info())
print("\\n--- Descriptive Statistics ---")
print(df.describe())

# Correlation Matrix (Numeric Only)
numeric_df = df.select_dtypes(include=[np.number])
if not numeric_df.empty:
    fig_corr = px.imshow(numeric_df.corr(), text_auto=True, title="Feature Correlation Heatmap", color_continuous_scale='RdBu_r')
    fig_corr.show()

# ==========================================
# 3. MACHINE LEARNING PIPELINE
# ==========================================
print("\\n--- Initializing ML Pipeline ---")

# A. Detect Target (Assuming last column for template)
target_col = df.columns[-1] 
print(f"🎯 Target Variable: {{target_col}}")

# B. Preprocessing
X = df.drop(columns=[target_col])
y = df[target_col]

# Handle Missing Values
imputer = SimpleImputer(strategy='mean')
feat_names = X.columns
X_imputed = pd.DataFrame(imputer.fit_transform(X.select_dtypes(include=[np.number])), columns=X.select_dtypes(include=[np.number]).columns)

# Handle Categorical Encoding
# (Simple Label Encoding for Tree Models)
for col in X.select_dtypes(include=['object', 'category']).columns:
    le = LabelEncoder()
    # Fill NA prevents crash
    X[col] = X[col].fillna('Unknown')
    X[col] = le.fit_transform(X[col])
    X_imputed[col] = X[col] # Add back to imputed

X_final = X_imputed
    
# C. Train/Test Split
X_train, X_test, y_train, y_test = train_test_split(X_final, y, test_size=0.2, random_state=42)

# D. Model Selection & Training
# Check if Regression or Classification based on target Type
is_regression = np.issubdtype(y.dtype, np.number) and y.nunique() > 20

if is_regression:
    model = RandomForestRegressor(n_estimators=100, random_state=42)
    print("🚀 Training Random Forest Regressor...")
else:
    model = RandomForestClassifier(n_estimators=100, random_state=42)
    print("🚀 Training Random Forest Classifier...")

model.fit(X_train, y_train)
score = model.score(X_test, y_test)
print(f"✅ Model Trained. Test Score (R2/Accuracy): {{score:.4f}}")

# ==========================================
# 4. EXPLAINABILITY (SHAP)
# ==========================================
print("\\n--- Generating Explanations ---")
# SHAP requires the raw model and the data
explainer = shap.Explainer(model)
shap_values = explainer(X_test)

# Plot 1: Summary Beeswarm
plt.figure()
shap.plots.beeswarm(shap_values, show=True)

# Plot 2: Bar Importance
plt.figure()
shap.plots.bar(shap_values, show=True)

print("✅ Analysis Complete.")
"""
                buffer = BytesIO()
                buffer.write(py_code.encode())
                buffer.seek(0)
                st.session_state['last_py'] = buffer
                st.rerun()
            
            if 'last_py' in st.session_state:
                 st.download_button("📥 Download .py Script", st.session_state['last_py'], "analysis_export.py", "text/x-python", key="dl_py_ready", use_container_width=True)

        with col_adv_2:
            st.info("ℹ️ **Premium Insight**\n\nFor full **SHAP (SHapley Additive exPlanations)** values, please use the Python Export feature to run a deep dive analysis in a local Notebook environment. This ensures maximum performance for large datasets.")

        st.markdown("---")
        if st.button("🔄 Start New Report (Clear Cache)", type="secondary", use_container_width=True):
            keys = ['last_pdf', 'last_word', 'last_excel', 'last_pptx', 'last_py']
            for k in keys:
                if k in st.session_state: del st.session_state[k]
            st.rerun()
